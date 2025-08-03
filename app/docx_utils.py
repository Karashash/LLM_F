from typing import List, Tuple
from io import BytesIO
from docx import Document

def _cells(doc):
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                yield cell

def extract_blocks(source) -> Tuple[List[dict], Document]:
    doc = Document(source)
    blocks=[]
    for p in doc.paragraphs:
        blocks.append({"type":"p","object":p,"content":p.text})
    for cell in _cells(doc):
        blocks.append({"type":"cell","object":cell,
                       "content":"\n".join([p.text for p in cell.paragraphs])})
    return blocks, doc

def restore_blocks(doc:Document, blocks:List[dict], translated:List[str]):
    for blk, txt in zip(blocks, translated):
        if blk["type"]=="p":
            p=blk["object"]
            for r in p.runs: r.text=""
            p.text = txt
        else:
            cell = blk["object"]
            while len(cell.paragraphs)>1:
                cell._tc.remove(cell.paragraphs[0]._p)
            cell.paragraphs[0].text = txt
